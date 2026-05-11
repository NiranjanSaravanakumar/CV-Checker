import logging

logger = logging.getLogger(__name__)


def parse_resume(file_path: str, filename: str) -> str:
    """
    Dispatch to the correct parser based on file extension.
    Returns extracted plain text.
    """
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    parsers = {
        'pdf': _parse_pdf,
        'docx': _parse_docx,
        'txt': _parse_txt,
    }

    parser = parsers.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file extension: .{ext}")

    return parser(file_path)


# ---------------------------------------------------------------------------
# Individual parsers
# ---------------------------------------------------------------------------

def _parse_pdf(file_path: str) -> str:
    """Extract text from a PDF using PyPDF2."""
    try:
        import PyPDF2
        pages = []
        with open(file_path, 'rb') as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return '\n'.join(pages)
    except Exception as exc:
        logger.error("PDF parsing error: %s", exc)
        raise


def _parse_docx(file_path: str) -> str:
    """Extract text from a DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return '\n'.join(paragraphs)
    except Exception as exc:
        logger.error("DOCX parsing error: %s", exc)
        raise


def _parse_txt(file_path: str) -> str:
    """Read a plain-text file, trying UTF-8 then falling back to latin-1."""
    for encoding in ('utf-8', 'latin-1', 'cp1252'):
        try:
            with open(file_path, 'r', encoding=encoding) as fh:
                return fh.read()
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file with any supported encoding")
